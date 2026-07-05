from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def iter_block_text(doc: Document) -> list[str]:
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                parts.append(" | ".join(cells))
    return parts


def main() -> int:
    if len(sys.argv) < 3:
        print("usage: extract_docx_questions.py OUT_JSON DOCX...", file=sys.stderr)
        return 2

    out_path = Path(sys.argv[1])
    records = []

    for raw_path in sys.argv[2:]:
        path = Path(raw_path)
        doc = Document(path)
        text_blocks = iter_block_text(doc)
        records.append(
            {
                "file": str(path),
                "name": path.name,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "blocks": text_blocks,
                "text": "\n".join(text_blocks),
            }
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
